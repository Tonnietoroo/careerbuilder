from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse, FileResponse
from .models import CV, CoverLetter, DocumentFolder, Document, CVTemplate
from .forms import CVForm, CoverLetterForm, InterviewPrepForm
from .utils import (
    generate_pdf, 
    generate_docx, 
    generate_cover_letter, 
    get_interview_questions, 
    generate_cover_letter_pdf, 
    generate_cover_letter_docx,
    generate_cover_letter_body,
    parse_cv_content,
    parse_cover_letter_content
)
from django.template.loader import get_template
import json
import openai
from django.core.paginator import Paginator
from django.db.models import Q
from interviews.models import InterviewPrep, InterviewQuestion
from django.utils import timezone
from django.core.files.storage import default_storage
import docx2txt
import PyPDF2
import re
from django.template.loader import render_to_string
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import os
import pdfkit
from django.conf import settings
import subprocess
import logging
from django.contrib.staticfiles import finders
from documents.models import Document
from .choices import CERTIFICATION_CHOICES


logger = logging.getLogger(__name__)

@login_required
def cv_list(request):
    cvs = CV.objects.filter(user=request.user).order_by('-updated_at')
    return render(request, 'documents/cv_list.html', {'cvs': cvs})

@login_required
def cv_template_select(request):
    templates = CVTemplate.objects.filter(is_active=True)
    return render(request, 'documents/cv_template_select.html', {
        'templates': templates
    })

@login_required
def cv_create(request):
    # Get template_id from either GET (initial) or POST (form submission)
    template_id = request.POST.get('template') or request.GET.get('template')
    if not template_id:
        return redirect('documents:cv_template_select')
    
    if request.method == 'POST':
        form = CVForm(request.POST)
        if form.is_valid():
            cv = form.save(commit=False)
            cv.user = request.user
            cv.template = get_object_or_404(CVTemplate, id=template_id)
            
            # Process the form data into a structured format
            content = {
                'personal_info': {
                    'full_name': form.cleaned_data['full_name'],
                    'email': form.cleaned_data['email'],
                    'phone_number': form.cleaned_data['phone_number'],
                    'address': form.cleaned_data['address'],
                    'gender': form.cleaned_data['gender'],
                    'date_of_birth': form.cleaned_data['date_of_birth'].strftime('%Y-%m-%d'),
                    'nationality': form.cleaned_data['nationality']
                },
                'personal_profile': form.cleaned_data['personal_profile'],
                'work_experience': [],
                'education': [],
                'skills': [],
                'certifications': [],
                'languages': [],
                'hobbies': [],
                'referees': []
            }
            
            # Process work experience
            work_experience = zip(
                request.POST.getlist('job_title[]'),
                request.POST.getlist('company_name[]'),
                request.POST.getlist('location[]'),
                request.POST.getlist('start_date[]'),
                request.POST.getlist('end_date[]'),
                request.POST.getlist('responsibilities[]')
            )
            
            for job in work_experience:
                content['work_experience'].append({
                    'job_title': job[0],
                    'company_name': job[1],
                    'location': job[2],
                    'start_date': job[3],
                    'end_date': job[4] if job[4] else None,
                    'responsibilities': job[5].split('\n')
                })
            
            # Process education
            education = zip(
                request.POST.getlist('degree[]'),
                request.POST.getlist('institution[]'),
                request.POST.getlist('location[]'),
                request.POST.getlist('start_date[]'),
                request.POST.getlist('end_date[]'),
                request.POST.getlist('grade[]')
            )
            
            for edu in education:
                content['education'].append({
                    'degree': edu[0],
                    'institution': edu[1],
                    'location': edu[2],
                    'start_date': edu[3],
                    'end_date': edu[4] if edu[4] else None,
                    'grade': edu[5]
                })
            
            # Process skills
            skills = zip(
                request.POST.getlist('skill_name[]'),
                request.POST.getlist('skill_type[]')
            )
            
            for skill in skills:
                content['skills'].append({
                    'name': skill[0],
                    'type': skill[1]
                })
            
            # Process certifications with new format
            certifications = []
            cert_choices = request.POST.getlist('certifications[][certificate_choice]')
            custom_certs = request.POST.getlist('certifications[][custom_certificate]')
            organizations = request.POST.getlist('certifications[][issuing_organization]')
            dates = request.POST.getlist('certifications[][date_earned]')
            
            for i in range(len(cert_choices)):
                cert_name = custom_certs[i] if cert_choices[i] == 'other' else cert_choices[i]
                certifications.append({
                    'certificate_name': cert_name,
                    'issuing_organization': organizations[i],
                    'date_earned': dates[i]
                })
            
            content['certifications'] = certifications
            
            # Process languages
            languages = zip(
                request.POST.getlist('language[]'),
                request.POST.getlist('proficiency[]')
            )
            
            for lang in languages:
                content['languages'].append({
                    'language': lang[0],
                    'proficiency': lang[1]
                })
            
            # Process hobbies
            content['hobbies'] = request.POST.getlist('hobby[]')
            
            # Process referees
            referees = zip(
                request.POST.getlist('referee_name[]'),
                request.POST.getlist('referee_job[]'),
                request.POST.getlist('referee_email[]'),
                request.POST.getlist('referee_phone[]')
            )
            
            for ref in referees:
                if any(ref):  # Only add if at least one field is filled
                    content['referees'].append({
                        'name': ref[0],
                        'job': ref[1],
                        'email': ref[2],
                        'phone': ref[3]
                    })
            
            cv.content = content
            cv.save()
            
            messages.success(request, 'CV created successfully!')
            return redirect('documents:cv_preview', cv_id=cv.id)
    else:
        form = CVForm()
    
    return render(request, 'documents/cv_create.html', {
        'form': form,
        'template_id': template_id,
        'certification_choices': CERTIFICATION_CHOICES,
        'edit_mode': False
    })

@login_required
def cv_preview(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id, user=request.user)
    
    # Get template CSS
    template_css = ""
    if cv.template and cv.template.css_template:
        try:
            template_css = cv.template.css_template.read().decode('utf-8')
        except Exception as e:
            messages.warning(request, "Could not load template CSS. Using default styling.")
    
    # Get template name (lowercase)
    template_name = cv.template.name.lower() if cv.template else 'professional'
    
    context = {
        'cv': cv,
        'template_css': template_css,
        'template_name': template_name,
        'personal_info': cv.content['personal_info'],
        'bio_data': cv.content.get('bio_data', {}),
        'personal_profile': cv.content['personal_profile'],
        'work_experience': cv.content['work_experience'],
        'education': cv.content['education'],
        'skills': cv.content['skills'],
        'certifications': cv.content.get('certifications', []),
        'languages': cv.content.get('languages', []),
        'hobbies': cv.content.get('hobbies', []),
        'referees': cv.content.get('referees', [])
    }
    
    return render(request, 'documents/cv_preview.html', context)

@login_required
def cv_edit(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id, user=request.user)
    if request.method == 'POST':
        form = CVForm(request.POST, instance=cv)
        if form.is_valid():
            cv = form.save(commit=False)
            content = cv.content
            
            # Update personal info while preserving existing data
            content['personal_info'].update({
                'full_name': form.cleaned_data['full_name'],
                'email': form.cleaned_data['email'],
                'phone_number': form.cleaned_data['phone_number'],
                'address': form.cleaned_data['address'],
                'gender': form.cleaned_data['gender'],
                'date_of_birth': form.cleaned_data['date_of_birth'].strftime('%Y-%m-%d'),
                'nationality': form.cleaned_data['nationality']
            })
            
            # Update personal profile
            content['personal_profile'] = form.cleaned_data['personal_profile']
            
            # Process work experience
            if request.POST.getlist('job_title[]'):
                content['work_experience'] = []
                work_experience = zip(
                    request.POST.getlist('job_title[]'),
                    request.POST.getlist('company_name[]'),
                    request.POST.getlist('location[]'),
                    request.POST.getlist('start_date[]'),
                    request.POST.getlist('end_date[]'),
                    request.POST.getlist('responsibilities[]')
                )
                
                for job in work_experience:
                    content['work_experience'].append({
                        'job_title': job[0],
                        'company_name': job[1],
                        'location': job[2],
                        'start_date': job[3],
                        'end_date': job[4] if job[4] else None,
                        'responsibilities': job[5].split('\n')
                    })
            
            # Process education
            if request.POST.getlist('degree[]'):
                content['education'] = []
                education = zip(
                    request.POST.getlist('degree[]'),
                    request.POST.getlist('institution[]'),
                    request.POST.getlist('location[]'),
                    request.POST.getlist('start_date[]'),
                    request.POST.getlist('end_date[]'),
                    request.POST.getlist('grade[]')
                )
                
                for edu in education:
                    content['education'].append({
                        'degree': edu[0],
                        'institution': edu[1],
                        'location': edu[2],
                        'start_date': edu[3],
                        'end_date': edu[4] if edu[4] else None,
                        'grade': edu[5]
                    })
            
            # Process skills
            if request.POST.getlist('skill_name[]'):
                content['skills'] = []
                skills = zip(
                    request.POST.getlist('skill_name[]'),
                    request.POST.getlist('skill_type[]')
                )
                
                for skill in skills:
                    content['skills'].append({
                        'name': skill[0],
                        'type': skill[1]
                    })
            
            # Process certifications
            if request.POST.getlist('certifications[][certificate_choice]'):
                content['certifications'] = []
                cert_choices = request.POST.getlist('certifications[][certificate_choice]')
                custom_certs = request.POST.getlist('certifications[][custom_certificate]')
                organizations = request.POST.getlist('certifications[][issuing_organization]')
                dates = request.POST.getlist('certifications[][date_earned]')
                
                for i in range(len(cert_choices)):
                    cert_name = custom_certs[i] if cert_choices[i] == 'other' else cert_choices[i]
                    content['certifications'].append({
                        'certificate_name': cert_name,
                        'issuing_organization': organizations[i],
                        'date_earned': dates[i]
                    })
            
            # Process languages
            if request.POST.getlist('language[]'):
                content['languages'] = []
                languages = zip(
                    request.POST.getlist('language[]'),
                    request.POST.getlist('proficiency[]')
                )
                
                for lang in languages:
                    content['languages'].append({
                        'language': lang[0],
                        'proficiency': lang[1]
                    })
            
            # Process hobbies
            if request.POST.getlist('hobby[]'):
                content['hobbies'] = request.POST.getlist('hobby[]')
            
            # Process referees
            if request.POST.getlist('referee_name[]'):
                content['referees'] = []
                referees = zip(
                    request.POST.getlist('referee_name[]'),
                    request.POST.getlist('referee_job[]'),
                    request.POST.getlist('referee_email[]'),
                    request.POST.getlist('referee_phone[]')
                )
                
                for ref in referees:
                    if any(ref):  # Only add if at least one field is filled
                        content['referees'].append({
                            'name': ref[0],
                            'job': ref[1],
                            'email': ref[2],
                            'phone': ref[3]
                        })
            
            cv.content = content
            cv.save()
            
            messages.success(request, 'CV updated successfully!')
            return redirect('documents:cv_preview', cv_id=cv.id)
    else:
        # Pre-fill the form with existing data
        personal_info = cv.content.get('personal_info', {})
        initial_data = {
            'title': cv.title,
            'full_name': personal_info.get('full_name', ''),
            'email': personal_info.get('email', ''),
            'phone_number': personal_info.get('phone_number', ''),
            'address': personal_info.get('address', ''),
            'gender': personal_info.get('gender', 'M'),  # Default to 'M' if not set
            'date_of_birth': personal_info.get('date_of_birth', ''),
            'nationality': personal_info.get('nationality', ''),
            'personal_profile': cv.content.get('personal_profile', '')
        }
        form = CVForm(initial=initial_data)
    
    return render(request, 'documents/cv_create.html', {
        'form': form,
        'cv': cv,
        'certification_choices': CERTIFICATION_CHOICES,
        'edit_mode': True,
        'work_experience': cv.content.get('work_experience', []),
        'education': cv.content.get('education', []),
        'skills': cv.content.get('skills', []),
        'certifications': cv.content.get('certifications', []),
        'languages': cv.content.get('languages', []),
        'hobbies': cv.content.get('hobbies', []),
        'referees': cv.content.get('referees', [])
    })

@login_required
def cv_delete(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id, user=request.user)
    if request.method == 'POST':
        cv.delete()
        messages.success(request, 'CV deleted successfully!')
    return redirect('documents:cv_list')

@login_required
def cv_download(request, cv_id):
    cv = get_object_or_404(CV, id=cv_id, user=request.user)
    download_format = request.GET.get('format', 'pdf')
    template_name = cv.template.name.lower() if cv.template else 'professional'

    if download_format == 'pdf':
        try:
            # Create a temporary HTML file with embedded CSS
            template = get_template(f'documents/templates/{template_name}_cv.html')
            
            # Read the CSS file using Django's static file finder
            css_path = finders.find('documents/css/cv_templates.css')
            if not css_path:
                raise FileNotFoundError("CSS file not found in static files")
                
            with open(css_path, 'r') as css_file:
                css_content = css_file.read()
            
            # Render the template with content and embedded CSS
            html_content = template.render({
                'cv': cv,
                'download_mode': True,
                'embedded_css': css_content,
                'personal_info': cv.content.get('personal_info', {}),
                'bio_data': cv.content.get('bio_data', {}),
                'personal_profile': cv.content.get('personal_profile', ''),
                'work_experience': cv.content.get('work_experience', []),
                'education': cv.content.get('education', []),
                'skills': cv.content.get('skills', []),
                'certifications': cv.content.get('certifications', []),
                'languages': cv.content.get('languages', []),
                'hobbies': cv.content.get('hobbies', []),
                'referees': cv.content.get('referees', [])
            })
            
            # Add required meta tags and CSS with improved font loading and print styles
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{cv.title}</title>
                <link href="https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600&family=Playfair+Display:wght@400;500;600;700&display=swap" rel="stylesheet">
                <style>
                    {css_content}
                    /* Enhanced print-specific CSS */
                    @page {{
                        margin: 0;
                        size: A4;
                    }}
                    body {{
                        margin: 0;
                        padding: 0;
                        background: white;
                        -webkit-print-color-adjust: exact !important;
                        print-color-adjust: exact !important;
                    }}
                    .cv-container {{
                        padding: 40px !important;
                        max-width: none !important;
                        width: 100% !important;
                        margin: 0 !important;
                    }}
                    /* Ensure sidebar prints correctly for creative template */
                    .creative .sidebar {{
                        position: fixed !important;
                        print-color-adjust: exact !important;
                        -webkit-print-color-adjust: exact !important;
                    }}
                    /* Ensure skill bars print correctly */
                    .skill-bar, .language-bar {{
                        print-color-adjust: exact !important;
                        -webkit-print-color-adjust: exact !important;
                    }}
                    /* Ensure proper page breaks */
                    .section {{
                        page-break-inside: avoid;
                    }}
                    .experience-item, .education-item {{
                        page-break-inside: avoid;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Create temp directory if it doesn't exist
            temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp')
            os.makedirs(temp_dir, exist_ok=True)
            
            # Use secure temporary file names
            import uuid
            temp_file_name = str(uuid.uuid4())
            temp_html = os.path.join(temp_dir, f'{temp_file_name}.html')
            temp_pdf = os.path.join(temp_dir, f'{temp_file_name}.pdf')
            
            # Save HTML content
            with open(temp_html, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Enhanced wkhtmltopdf options for better rendering
            options = [
                '--enable-local-file-access',
                '--encoding utf-8',
                '--javascript-delay 2000',
                '--no-stop-slow-scripts',
                '--enable-external-links',
                '--enable-internal-links',
                '--margin-top 0',
                '--margin-right 0',
                '--margin-bottom 0',
                '--margin-left 0',
                '--page-size A4',
                '--dpi 300',
                '--image-dpi 300',
                '--image-quality 100',
                '--print-media-type',
                '--enable-smart-shrinking',
            ]
            
            cmd = f'"{settings.WKHTMLTOPDF_PATH}" {" ".join(options)} "{temp_html}" "{temp_pdf}"'
            subprocess.run(cmd, shell=True, check=True)
            
            # Read the PDF and create response
            with open(temp_pdf, 'rb') as f:
                pdf_content = f.read()
                
            # Clean up temporary files
            try:
                os.remove(temp_html)
                os.remove(temp_pdf)
            except OSError:
                pass  # Ignore cleanup errors
            
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{cv.title}.pdf"'
            return response
            
        except Exception as e:
            logger.error(f'PDF generation failed for CV {cv_id}: {str(e)}')
            messages.error(request, f'PDF generation failed: {str(e)}. Please try downloading as Word document.')
            return redirect('documents:cv_preview', cv_id=cv.id)
    
    elif download_format == 'docx':
        try:
            # Create document and initialize styles
            doc = DocxDocument()
            styles = doc.styles
            
            # Create custom styles
            title_style = styles.add_style('CV Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            
            heading_style = styles.add_style('CV Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 139)
            
            normal_style = styles.add_style('CV Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.size = Pt(11)
            
            # Get CV content from JSONField
            content = cv.content
            personal_info = content.get('personal_info', {})
            
            # Add personal information
            doc.add_paragraph(personal_info.get('full_name', ''), style='CV Title')
            contact_info = doc.add_paragraph(style='CV Normal')
            contact_info.add_run(f"Email: {personal_info.get('email', '')}\n")
            contact_info.add_run(f"Phone: {personal_info.get('phone_number', '')}\n")
            contact_info.add_run(f"Address: {personal_info.get('address', '')}")
            
            # Add professional profile
            doc.add_paragraph('Professional Profile', style='CV Heading')
            doc.add_paragraph(content.get('personal_profile', ''), style='CV Normal')
            
            # Add work experience
            doc.add_paragraph('Work Experience', style='CV Heading')
            for exp in content.get('work_experience', []):
                p = doc.add_paragraph(style='CV Normal')
                p.add_run(f"{exp.get('job_title', '')} at {exp.get('company_name', '')}\n").bold = True
                p.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}\n").italic = True
                
                # Add responsibilities
                for resp in exp.get('responsibilities', []):
                    doc.add_paragraph(resp, style='List Bullet')
            
            # Add education
            doc.add_paragraph('Education', style='CV Heading')
            for edu in content.get('education', []):
                p = doc.add_paragraph(style='CV Normal')
                p.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}\n").bold = True
                p.add_run(f"{edu.get('start_date', '')} - {edu.get('end_date', 'Present')}\n").italic = True
                if edu.get('grade'):
                    p.add_run(f"Grade: {edu.get('grade', '')}")
            
            # Add skills
            if content.get('skills'):
                doc.add_paragraph('Skills', style='CV Heading')
                skills_para = doc.add_paragraph(style='CV Normal')
                skills_text = ', '.join(skill.get('name', '') for skill in content.get('skills', []))
                skills_para.add_run(skills_text)
            
            # Add certifications
            if content.get('certifications'):
                doc.add_paragraph('Certifications', style='CV Heading')
                for cert in content.get('certifications', []):
                    p = doc.add_paragraph(style='CV Normal')
                    p.add_run(f"{cert.get('name', '')} - {cert.get('organization', '')}")
                    if cert.get('date_earned'):
                        p.add_run(f" ({cert.get('date_earned', '')})").italic = True
            
            # Add languages
            if content.get('languages'):
                doc.add_paragraph('Languages', style='CV Heading')
                for lang in content.get('languages', []):
                    doc.add_paragraph(
                        f"{lang.get('language', '')} - {lang.get('proficiency', '')}",
                        style='CV Normal'
                    )
            
            # Add hobbies
            if content.get('hobbies'):
                doc.add_paragraph('Hobbies & Interests', style='CV Heading')
                hobbies_para = doc.add_paragraph(style='CV Normal')
                hobbies_text = ', '.join(hobby for hobby in content.get('hobbies', []))
                hobbies_para.add_run(hobbies_text)
            
            # Add referees
            if content.get('referees'):
                doc.add_paragraph('References', style='CV Heading')
                for ref in content.get('referees', []):
                    p = doc.add_paragraph(style='CV Normal')
                    p.add_run(f"{ref.get('name', '')} - {ref.get('job', '')}\n").bold = True
                    p.add_run(f"Email: {ref.get('email', '')}\n")
                    p.add_run(f"Phone: {ref.get('phone', '')}")
            
            # Set up response
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{cv.title}.docx"'
            doc.save(response)
            return response
            
        except Exception as e:
            logger.error(f'Word document generation failed for CV {cv_id}: {str(e)}')
            messages.error(request, f'Word document generation failed: {str(e)}')
            return redirect('documents:cv_preview', cv_id=cv.id)

    return redirect('documents:cv_preview', cv_id=cv.id)

def apply_professional_style(doc, cv):
    # Professional template styling
    content = cv.content
    
    # Header with name and contact info
    heading = doc.add_paragraph(content['personal_info']['full_name'], style='CV Title')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact information
    contact_para = doc.add_paragraph(style='CV Normal')
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(
        f"{content['personal_info']['email']} | {content['personal_info']['phone_number']}\n"
        f"{content['personal_info']['address']}"
    )
    
    # Professional Profile
    doc.add_paragraph('Professional Profile', style='CV Heading 1')
    doc.add_paragraph(content['personal_profile'], style='CV Normal')
    
    # Work Experience
    doc.add_paragraph('Work Experience', style='CV Heading 1')
    for exp in content['work_experience']:
        p = doc.add_paragraph(style='CV Normal')
        p.add_run(f"{exp['job_title']} - {exp['company_name']}").bold = True
        p.add_run(f"\n{exp['start_date']} - {exp['end_date'] or 'Present'}").italic = True
        for resp in exp['responsibilities']:
            bullet_p = doc.add_paragraph(style='List Bullet')
            bullet_p.add_run(resp)
    
    # Education
    doc.add_paragraph('Education', style='CV Heading 1')
    for edu in content['education']:
        p = doc.add_paragraph(style='CV Normal')
        p.add_run(f"{edu['degree']} - {edu['institution']}").bold = True
        p.add_run(f"\n{edu['start_date']} - {edu['end_date'] or 'Present'}").italic = True
        if edu.get('grade'):
            p.add_run(f"\nGrade: {edu['grade']}")
    
    # Skills
    if content.get('skills'):
        doc.add_paragraph('Skills', style='CV Heading 1')
        skills_para = doc.add_paragraph(style='CV Normal')
        for i, skill in enumerate(content['skills']):
            skills_para.add_run(skill['name'])
            if i < len(content['skills']) - 1:
                skills_para.add_run(' • ')

def apply_minimal_style(doc, cv):
    # Minimal template styling
    content = cv.content
    
    # Simple header
    heading = doc.add_paragraph(content['personal_info']['full_name'], style='CV Title')
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Contact info - clean and simple
    contact_para = doc.add_paragraph(style='CV Normal')
    contact_para.add_run(
        f"{content['personal_info']['email']} • "
        f"{content['personal_info']['phone_number']} • "
        f"{content['personal_info']['address']}"
    )
    
    # Profile
    doc.add_paragraph('Profile', style='CV Heading 1')
    doc.add_paragraph(content['personal_profile'], style='CV Normal')
    
    # Experience
    doc.add_paragraph('Experience', style='CV Heading 1')
    for exp in content['work_experience']:
        p = doc.add_paragraph(style='CV Normal')
        p.add_run(f"{exp['job_title']} | {exp['company_name']}").bold = True
        p.add_run(f"\n{exp['start_date']} - {exp['end_date'] or 'Present'}").italic = True
        for resp in exp['responsibilities']:
            bullet_p = doc.add_paragraph(style='List Bullet')
            bullet_p.add_run(resp)
    
    # Education
    doc.add_paragraph('Education', style='CV Heading 1')
    for edu in content['education']:
        p = doc.add_paragraph(style='CV Normal')
        p.add_run(f"{edu['degree']}\n{edu['institution']}").bold = True
        p.add_run(f"\n{edu['start_date']} - {edu['end_date'] or 'Present'}").italic = True
        if edu.get('grade'):
            p.add_run(f"\nGrade: {edu['grade']}")
    
    # Skills - minimal layout
    if content.get('skills'):
        doc.add_paragraph('Skills', style='CV Heading 1')
        skills_text = ', '.join(skill['name'] for skill in content['skills'])
        doc.add_paragraph(skills_text, style='CV Normal')
            
def apply_creative_style(doc, cv):
    # Creative template styling
    content = cv.content
    
    # Stylized header
    heading = doc.add_paragraph(content['personal_info']['full_name'], style='CV Title')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Creative contact layout
    contact_para = doc.add_paragraph(style='CV Normal')
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(
        f"📧 {content['personal_info']['email']} | "
        f"📱 {content['personal_info']['phone_number']}\n"
        f"📍 {content['personal_info']['address']}"
    )
    
    # About Me section
    doc.add_paragraph('About Me', style='CV Heading 1')
    profile_para = doc.add_paragraph(style='CV Normal')
    profile_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    profile_para.add_run(content['personal_profile'])
    
    # Experience with creative formatting
    doc.add_paragraph('Professional Journey', style='CV Heading 1')
    for exp in content['work_experience']:
        p = doc.add_paragraph(style='CV Normal')
        p.add_run('🎯 ').font.size = Pt(14)
        p.add_run(f"{exp['job_title']} - {exp['company_name']}").bold = True
        p.add_run(f"\n📅 {exp['start_date']} - {exp['end_date'] or 'Present'}").italic = True
        for resp in exp['responsibilities']:
            bullet_p = doc.add_paragraph(style='CV Normal')
            bullet_p.add_run('✓ ' + resp)
    
    # Education with creative elements
    doc.add_paragraph('Educational Background', style='CV Heading 1')
    for edu in content['education']:
        p = doc.add_paragraph(style='CV Normal')
        p.add_run('🎓 ').font.size = Pt(14)
        p.add_run(f"{edu['degree']}\n{edu['institution']}").bold = True
        p.add_run(f"\n📅 {exp['start_date']} - {exp['end_date'] or 'Present'}").italic = True
        if edu.get('grade'):
            p.add_run(f"\n📊 Grade: {edu['grade']}")
    
    # Skills with creative layout
    if content.get('skills'):
        doc.add_paragraph('Skills & Expertise', style='CV Heading 1')
        skills_para = doc.add_paragraph(style='CV Normal')
        for i, skill in enumerate(content['skills']):
            skills_para.add_run('💫 ' + skill['name'])
            if i < len(content['skills']) - 1:
                skills_para.add_run(' | ')

@login_required
def cover_letter_create(request):
    if request.method == 'POST':
        form = CoverLetterForm(request.user, request.POST)
        if form.is_valid():
            cover_letter = form.save(commit=False)
            cover_letter.user = request.user
            
            try:
                # Get CV data for the selected CV
                cv = form.cleaned_data['cv']
                personal_info = cv.content.get('personal_info', {})
                
                # Format the current date
                current_date = timezone.now().strftime("%B %d, %Y")
                
                # Generate the cover letter body
                body = generate_cover_letter_body(
                    cv=cv,
                    job_title=form.cleaned_data['job_title'],
                    company_name=form.cleaned_data['company_name']
                )
                
                # Structure the cover letter content
                content = {
                    'company_name': form.cleaned_data['company_name'],
                    'job_title': form.cleaned_data['job_title'],
                    'personal_info': {
                        'full_name': personal_info.get('full_name', ''),
                        'address': personal_info.get('address', ''),
                        'phone': personal_info.get('phone_number', ''),
                        'email': personal_info.get('email', '')
                    },
                    'date': current_date,
                    'body': body
                }
                
                # Store content as JSON string
                cover_letter.content = json.dumps(content)
                cover_letter.save()
                
                messages.success(request, 'Cover letter created successfully!')
                return redirect('documents:cover_letter_preview', cover_letter_id=cover_letter.id)
                
            except Exception as e:
                messages.error(request, f'Error generating cover letter: {str(e)}')
                return redirect('documents:cover_letter_list')
    else:
        form = CoverLetterForm(request.user)
    
    return render(request, 'documents/cover_letter_create.html', {'form': form})

@login_required
def cover_letter_preview(request, cover_letter_id):
    cover_letter = get_object_or_404(CoverLetter, id=cover_letter_id, user=request.user)
    
    try:
        # Convert string content to dictionary if needed
        content = cover_letter.content
        if isinstance(content, str):
            content = json.loads(content)
            
        # Ensure all required content is present
        if not content.get('body'):
            messages.error(request, 'Error: Cover letter content is missing')
            return redirect('documents:cover_letter_list')
        
        # Update cover_letter object with parsed content
        cover_letter.content = content
        
    except (json.JSONDecodeError, AttributeError) as e:
        messages.error(request, 'Error: Invalid cover letter format')
        return redirect('documents:cover_letter_list')
    
    return render(request, 'documents/cover_letter_preview.html', {
        'cover_letter': cover_letter,
    })

@login_required
def cover_letter_list(request):
    cover_letters = CoverLetter.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'documents/cover_letter_list.html', {'cover_letters': cover_letters})

@login_required
def cover_letter_edit(request, cover_letter_id):
    cover_letter = get_object_or_404(CoverLetter, id=cover_letter_id, user=request.user)
    
    if request.method == 'POST':
        form = CoverLetterForm(request.user, request.POST, instance=cover_letter)
        if form.is_valid():
            cover_letter = form.save(commit=False)
            
            # Regenerate content if any field changed
            if (form.cleaned_data['cv'] != cover_letter.cv or 
                form.cleaned_data['job_title'] != cover_letter.job_title or
                form.cleaned_data['company_name'] != cover_letter.company_name or
                form.cleaned_data['company_address'] != cover_letter.company_address):
                
                try:
                    content = generate_cover_letter(
                        form.cleaned_data['cv'],
                        form.cleaned_data['job_title'],
                        form.cleaned_data['company_name']
                    )
                    cover_letter.content = content
                except Exception as e:
                    messages.error(request, 'Error regenerating cover letter content. Please try again.')
                    return render(request, 'documents/cover_letter_edit.html', {'form': form})
            
            cover_letter.save()
            messages.success(request, 'Cover letter updated successfully!')
            return redirect('documents:cover_letter_preview', cover_letter_id=cover_letter.id)
    else:
        form = CoverLetterForm(request.user, instance=cover_letter, initial={
            'cv': cover_letter.cv,
            'company_name': cover_letter.company_name,
            'company_address': cover_letter.company_address,
            'job_title': cover_letter.job_title
        })
    
    return render(request, 'documents/cover_letter_edit.html', {
        'form': form,
        'cover_letter': cover_letter
    })

@login_required
def interview_prep_create(request):
    if request.method == 'POST':
        form = InterviewPrepForm(request.POST)
        if form.is_valid():
            prep = form.save(commit=False)
            prep.user = request.user
            
            # Get interview questions using AI
            job_title = form.cleaned_data['job_title']
            company_name = form.cleaned_data['company_name']
            
            try:
                questions = get_interview_questions(job_title, company_name)
                prep.save()  # Save before adding many-to-many relationship
                prep.questions.set(questions)
                messages.success(request, 'Interview preparation created successfully!')
                return redirect('documents:interview_prep_detail', prep_id=prep.id)
            except Exception as e:
                messages.error(request, str(e))
                return render(request, 'documents/interview_prep.html', {'form': form})
    else:
        form = InterviewPrepForm()
    
    return render(request, 'documents/interview_prep.html', {'form': form})

@login_required
def interview_prep_detail(request, prep_id):
    prep = get_object_or_404(InterviewPrep, id=prep_id, user=request.user)
    questions = prep.questions.all()
    return render(request, 'documents/interview_prep_detail.html', {
        'prep': prep,
        'questions': questions
    })

@login_required
def interview_prep_list(request):
    preps = InterviewPrep.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'documents/interview_prep_list.html', {'preps': preps})

@login_required
def interview_prep_delete(request, prep_id):
    prep = get_object_or_404(InterviewPrep, id=prep_id, user=request.user)
    
    if request.method == 'POST':
        # Delete associated questions first
        prep.questions.all().delete()
        prep.delete()
        messages.success(request, 'Interview preparation deleted successfully!')
    
    return redirect('documents:interview_prep_list')

@login_required
def interview_prep_edit(request, prep_id):
    prep = get_object_or_404(InterviewPrep, id=prep_id, user=request.user)
    
    if request.method == 'POST':
        form = InterviewPrepForm(request.POST, instance=prep)
        if form.is_valid():
            form.save()
            messages.success(request, 'Interview preparation updated successfully!')
            return redirect('documents:interview_prep_detail', prep_id=prep.id)
    else:
        form = InterviewPrepForm(instance=prep)
    
    return render(request, 'documents/interview_prep.html', {'form': form, 'edit_mode': True})

@login_required
def document_manager(request):
    # Get all folders for the user
    folders = DocumentFolder.objects.filter(user=request.user)
    
    # Get all documents for the user, ordered by upload date
    documents = Document.objects.filter(user=request.user).order_by('-uploaded_at')
    
    context = {
        'folders': folders,
        'documents': documents,
        'current_folder': None
    }
    
    return render(request, 'documents/document_manager.html', context)

@login_required
def folder_view(request, folder_id):
    # Get the current folder
    folder = get_object_or_404(DocumentFolder, id=folder_id, user=request.user)
    
    # Get all folders for the sidebar
    folders = DocumentFolder.objects.filter(user=request.user)
    
    # Get documents in this folder
    documents = Document.objects.filter(
        user=request.user,
        folder=folder
    ).order_by('-uploaded_at')
    
    context = {
        'folders': folders,
        'documents': documents,
        'current_folder': folder
    }
    
    return render(request, 'documents/document_manager.html', context)

@login_required
def folder_create(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        description = request.POST.get('description')
        parent_id = request.POST.get('parent')
        
        try:
            folder = DocumentFolder(
                user=request.user,
                name=name,
                description=description
            )
            if parent_id:
                folder.parent = get_object_or_404(DocumentFolder, id=parent_id, user=request.user)
            folder.save()
            messages.success(request, 'Folder created successfully!')
        except Exception as e:
            messages.error(request, f'Error creating folder: {str(e)}')
    
    return redirect('documents:document_manager')

@login_required
def document_upload(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        document_type = request.POST.get('document_type')
        uploaded_file = request.FILES.get('file')
        
        if not all([title, document_type, uploaded_file]):
            messages.error(request, 'Please fill all required fields')
            return redirect('documents:document_manager')
        
        try:
            # Create document
            document = Document.objects.create(
                user=request.user,
                title=title,
                document_type=document_type,
                file=uploaded_file
            )
            
            # Extract content
            try:
                content = extract_document_content(document)
                if content:
                    document.content = content
                    document.is_parsed = True
                    document.save()
                    messages.success(request, 'Document uploaded and parsed successfully!')
                else:
                    messages.warning(request, 'Document uploaded but content could not be parsed')
            except ValueError as ve:
                messages.warning(request, str(ve))
            except Exception as e:
                logger.error(f"Error processing document {document.id}: {str(e)}")
                messages.error(request, f'Error processing document: {str(e)}')
            
            return redirect('documents:document_manager')
            
        except Exception as e:
            logger.error(f"Error uploading document: {str(e)}")
            messages.error(request, f'Error uploading document: {str(e)}')
            return redirect('documents:document_manager')
    
    return redirect('documents:document_manager')

def extract_document_content(document):
    """Extract content from uploaded document."""
    try:
        file_path = document.file.path
        text = ""
        
        # Extract text based on file type
        if file_path.lower().endswith('.docx'):
            try:
                text = docx2txt.process(file_path)
            except Exception as e:
                logger.error(f"Error processing DOCX file: {str(e)}")
                raise ValueError(f"Could not process DOCX file: {str(e)}")
                
        elif file_path.lower().endswith('.pdf'):
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
            except Exception as e:
                logger.error(f"Error processing PDF file: {str(e)}")
                raise ValueError(f"Could not process PDF file: {str(e)}")
                
        else:
            supported = ['.pdf', '.docx']
            raise ValueError(f"Unsupported file type. Supported types are: {', '.join(supported)}")
        
        if not text.strip():
            raise ValueError("No text could be extracted from the document")
        
        # Parse content based on document type
        content = None
        if document.document_type == 'cv':
            content = parse_cv_content(text)
            
            # Validate extracted content
            if not content['personal_info'].get('email') and not content['personal_info'].get('phone_number'):
                logger.warning(f"No contact information found in CV: {document.id}")
                
            if not content['work_experience'] and not content['education']:
                logger.warning(f"No work experience or education found in CV: {document.id}")
                
        elif document.document_type == 'cover_letter':
            content = parse_cover_letter_content(text)
        
        if not content:
            raise ValueError(f"Could not parse {document.document_type} content")
            
        return content
        
    except Exception as e:
        logger.error(f"Error extracting content from document {document.id}: {str(e)}")
        raise

@login_required
def convert_to_template(request, document_id):
    document = get_object_or_404(Document, id=document_id, user=request.user)
    
    if request.method == 'POST':
        template_id = request.POST.get('template')
        if not template_id:
            messages.error(request, 'Please select a template')
            return redirect('documents:document_manager')
        
        template = get_object_or_404(CVTemplate, id=template_id)
        
        try:
            # Extract content from the document if not already parsed
            if not document.is_parsed:
                content = extract_document_content(document)
                if content:
                    document.content = content
                    document.is_parsed = True
                    document.save()
                else:
                    raise ValueError("Could not parse document content")
            
            # Create CV content structure
            cv_content = {
                'personal_info': {
                    'full_name': document.content.get('personal_info', {}).get('full_name', ''),
                    'email': document.content.get('personal_info', {}).get('email', ''),
                    'phone_number': document.content.get('personal_info', {}).get('phone_number', ''),
                    'address': document.content.get('personal_info', {}).get('address', ''),
                    'gender': document.content.get('personal_info', {}).get('gender', 'M'),
                    'date_of_birth': document.content.get('personal_info', {}).get('date_of_birth', ''),
                    'nationality': document.content.get('personal_info', {}).get('nationality', '')
                },
                'personal_profile': document.content.get('personal_profile', ''),
                'work_experience': document.content.get('work_experience', []),
                'education': document.content.get('education', []),
                'skills': document.content.get('skills', []),
                'certifications': document.content.get('certifications', []),
                'languages': document.content.get('languages', []),
                'hobbies': document.content.get('hobbies', []),
                'referees': document.content.get('referees', [])
            }
            
            # Create new CV
            cv = CV.objects.create(
                user=request.user,
                title=f"{document.title} (Converted)",
                template=template,
                content=cv_content
            )
            
            messages.success(request, 'Document converted to CV successfully! You can now edit and customize it.')
            return redirect('documents:cv_edit', cv_id=cv.id)
            
        except Exception as e:
            messages.error(request, f'Error converting document: {str(e)}')
            return redirect('documents:document_manager')
    
    templates = CVTemplate.objects.filter(is_active=True)
    return render(request, 'documents/convert_to_template.html', {
        'document': document,
        'templates': templates
    })

@login_required
def document_move(request, document_id):
    document = get_object_or_404(Document, id=document_id, user=request.user)
    
    if request.method == 'POST':
        folder_id = request.POST.get('folder')
        try:
            if folder_id:
                document.folder = get_object_or_404(DocumentFolder, id=folder_id, user=request.user)
            else:
                document.folder = None
            document.save()
            messages.success(request, 'Document moved successfully!')
        except Exception as e:
            messages.error(request, f'Error moving document: {str(e)}')
    
    return redirect('documents:document_manager')

@login_required
def document_delete(request, document_id):
    if request.method == 'POST':
        document = get_object_or_404(Document, id=document_id, user=request.user)
        try:
            # Delete the actual file
            if document.file:
                document.file.delete()
            # Delete the document record
            document.delete()
            messages.success(request, 'Document deleted successfully')
        except Exception as e:
            messages.error(request, f'Error deleting document: {str(e)}')
    
    return redirect('documents:document_manager')

@login_required
def cover_letter_download(request, cover_letter_id):
    cover_letter = get_object_or_404(CoverLetter, id=cover_letter_id, user=request.user)
    format = request.GET.get('format', 'pdf')
    
    if format == 'pdf':
        buffer = generate_cover_letter_pdf(cover_letter)
        return FileResponse(
            buffer, 
            as_attachment=True, 
            filename=f'cover_letter_{cover_letter.company_name}.pdf'
        )
    elif format == 'docx':
        buffer = generate_cover_letter_docx(cover_letter)
        return FileResponse(
            buffer, 
            as_attachment=True, 
            filename=f'cover_letter_{cover_letter.company_name}.docx'
        )
    
    return HttpResponse('Invalid format specified')

@login_required
def cover_letter_delete(request, cover_letter_id):
    cover_letter = get_object_or_404(CoverLetter, id=cover_letter_id, user=request.user)
    
    if request.method == 'POST':
        cover_letter.delete()
        messages.success(request, 'Cover letter deleted successfully!')
    else:
        messages.error(request, 'Invalid request method.')
    
    return redirect('documents:cover_letter_list') 