from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
from typing import Dict
import io
import openai
from django.conf import settings
from interviews.models import InterviewQuestion
from io import BytesIO
from openai import OpenAI
import json
import re
import logging

def format_experience(experience):
    """Format work experience for the prompt"""
    formatted = []
    for job in experience:
        formatted.append(f"- {job['job_title']} at {job['company_name']}")
        for resp in job.get('responsibilities', []):
            formatted.append(f"  • {resp}")
    return "\n".join(formatted)

def format_skills(skills):
    """Format skills for the prompt"""
    hard_skills = [skill['name'] for skill in skills if skill['type'] == 'hard']
    soft_skills = [skill['name'] for skill in skills if skill['type'] == 'soft']
    
    formatted = []
    if hard_skills:
        formatted.append("Hard Skills:")
        formatted.extend([f"• {skill}" for skill in hard_skills])
    if soft_skills:
        formatted.append("\nSoft Skills:")
        formatted.extend([f"• {skill}" for skill in soft_skills])
    return "\n".join(formatted)

def generate_cover_letter(cv, job_title, company_name):
    """Generate a professional cover letter using OpenAI's GPT model"""
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=settings.OPENROUTER_API_KEY,
    )
    
    # Extract relevant information from CV
    personal_info = cv.content.get('personal_info', {})
    full_name = personal_info.get('full_name', '')
    email = personal_info.get('email', '')
    phone = personal_info.get('phone_number', '')
    address = personal_info.get('address', '')
    experience = cv.content.get('work_experience', [])
    skills = cv.content.get('skills', [])
    personal_profile = cv.content.get('personal_profile', '')
    
    from datetime import datetime
    current_date = datetime.now().strftime("%B %d, %Y")
    
    # Create professional template with proper spacing and formatting
    template = f"""
{full_name}
{address}
{phone}
{email}

{current_date}

Hiring Manager
{company_name}
[Company Address]

Dear Hiring Manager,

[CONTENT]

Best regards,
{full_name}"""

    # Create prompt for GPT with specific instructions for professional content
    prompt = f"""
    Generate three compelling paragraphs for a professional cover letter for a {job_title} position at {company_name}.
    
    Use these details strategically:
    - Skills: {format_skills(skills)}
    - Experience: {format_experience(experience)}
    - Profile: {personal_profile}

    Follow this structure:
    1. Opening Paragraph: 
       - Express genuine interest in the specific role
       - Demonstrate knowledge of {company_name}
       - Mention how your background aligns with the company's mission
       
    2. Core Value Paragraph:
       - Highlight 2-3 most relevant achievements
       - Connect your experience directly to the role's requirements
       - Use specific metrics or outcomes when possible
       
    3. Closing Paragraph:
       - Reinforce your enthusiasm
       - Summarize your value proposition
       - Express interest in further discussion
       
    Requirements:
    - Keep tone professional but engaging
    - Focus on value you'll bring to the role
    - Use active voice and confident language
    - Avoid clichés and generic statements
    - Keep paragraphs concise and impactful
    - Do not include any headers or contact information
    """
    
    try:
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": settings.SITE_URL,
                "X-Title": "CareerBuilder",
            },
            model="deepseek/deepseek-r1:free",
            messages=[
                {
                    "role": "system", 
                    "content": """You are an expert cover letter writer who creates compelling, professional content that:
                    - Demonstrates genuine interest in the role and company
                    - Highlights relevant achievements and skills
                    - Uses confident, active language
                    - Maintains professional tone while showing personality
                    - Creates clear value propositions
                    Generate only the body paragraphs, no headers or signatures."""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=800,
        )
        
        generated_content = completion.choices[0].message.content.strip()
        
        # Clean up any formatting issues
        content_lines = generated_content.split('\n')
        cleaned_lines = []
        for line in content_lines:
            # Skip any lines that might contain formatting markers or duplicated info
            if any(x in line.lower() for x in [
                'first paragraph:', 'second paragraph:', 'third paragraph:',
                'dear', 'sincerely', 'hiring manager', full_name.lower(), 
                email, phone, address, '[content]'
            ]):
                continue
            cleaned_lines.append(line)
        
        # Join lines and ensure proper spacing
        cleaned_content = '\n\n'.join(
            para.strip() 
            for para in ' '.join(cleaned_lines).split('\n\n') 
            if para.strip()
        )
        
        # Create the final letter with proper formatting
        final_letter = template.replace('[CONTENT]', cleaned_content)
        
        if len(final_letter) < 100:
            raise ValueError("Generated content too short")
            
        return final_letter
        
    except Exception as e:
        print(f"Error generating cover letter: {str(e)}")
        return template.replace('[CONTENT]', 
            f"I am writing to express my strong interest in the {job_title} position at {company_name}. "
            "With my background and skills, I believe I would be a valuable addition to your team.\n\n"
            "[Error: Could not generate custom content. Please try again.]"
        )

def generate_pdf(cv, output):
    """Generate a PDF version of the CV"""
    doc = SimpleDocTemplate(output, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Center alignment
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceBefore=20,
        spaceAfter=12,
        textColor=colors.HexColor('#2c3e50')
    )

    # Add name as title
    elements.append(Paragraph(cv.content['personal_info']['full_name'], title_style))
    
    # Contact information
    contact_info = [
        [cv.content['personal_info']['email']],
        [cv.content['personal_info']['phone_number']],
        [cv.content['personal_info']['address']]
    ]
    contact_table = Table(contact_info, colWidths=[400])
    contact_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#666666')),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    elements.append(contact_table)
    elements.append(Spacer(1, 20))

    # Professional Profile
    elements.append(Paragraph('Professional Profile', heading_style))
    elements.append(Paragraph(cv.content['personal_profile'], styles['Normal']))
    elements.append(Spacer(1, 20))

    # Work Experience
    elements.append(Paragraph('Work Experience', heading_style))
    for exp in cv.content['work_experience']:
        elements.append(Paragraph(
            f"<b>{exp['job_title']}</b> at <b>{exp['company_name']}</b>",
            styles['Normal']
        ))
        elements.append(Paragraph(
            f"<i>{exp['start_date']} - {exp['end_date'] or 'Present'}</i>",
            styles['Normal']
        ))
        for resp in exp['responsibilities']:
            elements.append(Paragraph(f"• {resp}", styles['Normal']))
        elements.append(Spacer(1, 12))

    # Education
    elements.append(Paragraph('Education', heading_style))
    for edu in cv.content['education']:
        elements.append(Paragraph(
            f"<b>{edu['degree']}</b> - {edu['institution']}",
            styles['Normal']
        ))
        elements.append(Paragraph(
            f"<i>{edu['start_date']} - {edu['end_date'] or 'Present'}</i>",
            styles['Normal']
        ))
        if edu['grade']:
            elements.append(Paragraph(f"Grade: {edu['grade']}", styles['Normal']))
        elements.append(Spacer(1, 12))

    # Skills
    if cv.content['skills']:
        elements.append(Paragraph('Skills', heading_style))
        skill_text = ' • '.join(skill['name'] for skill in cv.content['skills'])
        elements.append(Paragraph(skill_text, styles['Normal']))

    # Build the PDF
    doc.build(elements)

def generate_docx(cv, output):
    """Generate a Word document version of the CV"""
    doc = Document()
    
    # Add custom styles
    styles = doc.styles
    
    # Heading style
    heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
    heading_style.paragraph_format.space_before = Pt(20)
    heading_style.paragraph_format.space_after = Pt(12)
    
    # Name and title
    title = doc.add_heading(cv.content['personal_info']['full_name'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact information
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info.add_run(cv.content['personal_info']['email']).italic = True
    contact_info.add_run(' | ')
    contact_info.add_run(cv.content['personal_info']['phone_number']).italic = True
    contact_info.add_run(' | ')
    contact_info.add_run(cv.content['personal_info']['address']).italic = True
    
    # Professional Profile
    doc.add_heading('Professional Profile', level=1).style = heading_style
    doc.add_paragraph(cv.content['personal_profile'])
    
    # Work Experience
    doc.add_heading('Work Experience', level=1).style = heading_style
    for exp in cv.content['work_experience']:
        p = doc.add_paragraph()
        p.add_run(f"{exp['job_title']} at {exp['company_name']}").bold = True
        p.add_run(f"\n{exp['start_date']} - {exp['end_date'] or 'Present'}").italic = True
        
        if exp['responsibilities']:
            for resp in exp['responsibilities']:
                bullet_p = doc.add_paragraph(style='List Bullet')
                bullet_p.add_run(resp)
    
    # Education
    doc.add_heading('Education', level=1).style = heading_style
    for edu in cv.content['education']:
        p = doc.add_paragraph()
        p.add_run(f"{edu['degree']} - {edu['institution']}").bold = True
        p.add_run(f"\n{edu['start_date']} - {edu['end_date'] or 'Present'}").italic = True
        if edu['grade']:
            p.add_run(f"\nGrade: {edu['grade']}")
    
    # Skills
    if cv.content['skills']:
        doc.add_heading('Skills', level=1).style = heading_style
        skills_para = doc.add_paragraph()
        for i, skill in enumerate(cv.content['skills']):
            skills_para.add_run(skill['name'])
            if i < len(cv.content['skills']) - 1:
                skills_para.add_run(' • ')
    
    # Set margins for better layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Save the document
    doc.save(output)

def get_interview_questions(job_title, company_name):
    """Generate interview questions and tips for a specific job"""
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=settings.OPENROUTER_API_KEY,
        default_headers=settings.OPENROUTER_HEADERS
    )
    
    prompt = f"""
    Generate 10 interview questions and detailed answers for a {job_title} position at {company_name}.
    
    For each question, provide:
    1. A challenging but realistic question
    2. A detailed sample answer that demonstrates expertise
    3. Mix of technical, behavioral, and role-specific questions
    
    Format EXACTLY as follows for each question:
    Q: [The interview question]
    A: [The detailed answer]

    Make questions specific to the {job_title} role and {company_name}'s industry.
    Include questions about:
    - Technical skills required for the role
    - Past experience and achievements
    - Problem-solving abilities
    - Team collaboration
    - Role-specific scenarios
    """
    
    try:
        completion = client.chat.completions.create(
            model="deepseek/deepseek-r1:free",
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert interviewer. Generate specific, detailed interview questions and answers. Format each Q&A pair exactly as requested."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        content = completion.choices[0].message.content.strip()
        
        # Parse and store the questions
        questions = []
        current_question = None
        current_answer = []
        
        # Split content into lines and process
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('Q:'):
                # If we have a previous question, save it
                if current_question is not None:
                    question_obj = InterviewQuestion.objects.create(
                        question=current_question,
                        answer='\n'.join(current_answer),
                        question_type='general'
                    )
                    questions.append(question_obj)
                
                # Start new question
                current_question = line[2:].strip()
                current_answer = []
            elif line.startswith('A:'):
                current_answer.append(line[2:].strip())
            elif current_answer:  # Continue adding to current answer
                current_answer.append(line)
        
        # Don't forget to save the last question
        if current_question is not None:
            question_obj = InterviewQuestion.objects.create(
                question=current_question,
                answer='\n'.join(current_answer),
                question_type='general'
            )
            questions.append(question_obj)
        
        return questions
        
    except Exception as e:
        print(f"Error generating interview questions: {str(e)}")
        # Return default questions if generation fails
        return [
            InterviewQuestion.objects.create(
                question_type='general',
                question=f"Tell me about your experience relevant to this {job_title} position.",
                answer=f"Focus on your most relevant experience for the {job_title} role. Highlight specific achievements and how they demonstrate your qualifications."
            ),
            InterviewQuestion.objects.create(
                question_type='general',
                question=f"Why are you interested in working at {company_name}?",
                answer=f"Research {company_name}'s mission and values. Explain how your career goals align with the company's objectives."
            )
        ]

def generate_cover_letter_pdf(cover_letter):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Parse JSON content if it's a string
    content = cover_letter.content
    if isinstance(content, str):
        content = json.loads(content)

    # Add sender info
    elements.append(Paragraph(content['personal_info']['full_name'], styles['Heading1']))
    elements.append(Paragraph(content['personal_info']['address'], styles['Normal']))
    elements.append(Paragraph(content['personal_info']['phone'], styles['Normal']))
    elements.append(Paragraph(content['personal_info']['email'], styles['Normal']))
    elements.append(Spacer(1, 12))

    # Add date
    elements.append(Paragraph(content['date'], styles['Normal']))
    elements.append(Spacer(1, 12))

    # Add recipient info
    elements.append(Paragraph('Hiring Manager', styles['Normal']))
    elements.append(Paragraph(content['company_name'], styles['Normal']))
    if cover_letter.company_address:
        elements.append(Paragraph(cover_letter.company_address, styles['Normal']))
    elements.append(Spacer(1, 12))

    # Add salutation
    elements.append(Paragraph('Dear Hiring Manager,', styles['Normal']))
    elements.append(Spacer(1, 12))

    # Add body content
    body_paragraphs = content['body'].split('\n\n')
    for paragraph in body_paragraphs:
        if paragraph.strip():
            elements.append(Paragraph(paragraph.strip(), styles['Normal']))
            elements.append(Spacer(1, 12))

    # Add closing
    elements.append(Spacer(1, 12))
    elements.append(Paragraph('Best regards,', styles['Normal']))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph(content['personal_info']['full_name'], styles['Normal']))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_cover_letter_docx(cover_letter):
    doc = Document()
    
    # Parse JSON content if it's a string
    content = cover_letter.content
    if isinstance(content, str):
        content = json.loads(content)

    # Add sender info
    doc.add_paragraph(content['personal_info']['full_name'])
    doc.add_paragraph(content['personal_info']['address'])
    doc.add_paragraph(content['personal_info']['phone'])
    doc.add_paragraph(content['personal_info']['email'])
    doc.add_paragraph()  # blank line

    # Add date
    doc.add_paragraph(content['date'])
    doc.add_paragraph()  # blank line

    # Add recipient info
    doc.add_paragraph('Hiring Manager')
    doc.add_paragraph(content['company_name'])
    if cover_letter.company_address:
        doc.add_paragraph(cover_letter.company_address)
    doc.add_paragraph()  # blank line

    # Add salutation
    doc.add_paragraph('Dear Hiring Manager,')
    doc.add_paragraph()  # blank line

    # Add body content
    body_paragraphs = content['body'].split('\n\n')
    for paragraph in body_paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.add_paragraph()  # blank line

    # Add closing
    doc.add_paragraph('Best regards,')
    doc.add_paragraph(content['personal_info']['full_name'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_cover_letter_body(cv, job_title, company_name):
    """
    Generate the body content for a cover letter using CV data and job details.
    """
    try:
        # Get relevant information from CV
        work_experience = cv.content.get('work_experience', [])
        skills = cv.content.get('skills', [])
        education = cv.content.get('education', [])
        
        # Create an introduction paragraph
        intro = (f"I am writing to express my keen interest in the {job_title} position at {company_name}. "
                "As a motivated professional with a strong background in technology and a passion for innovation, "
                "I am confident in my ability to contribute effectively to your team.")

        # Create a skills and experience paragraph
        skill_names = [skill['name'] for skill in skills[:3]]  # Get top 3 skills
        if skill_names:
            skills_text = ", ".join(skill_names[:-1]) + f" and {skill_names[-1]}" if len(skill_names) > 1 else skill_names[0]
            experience_para = f"\n\nMy expertise in {skills_text} aligns perfectly with the requirements of this role. "
        else:
            experience_para = "\n\nThroughout my career, I have developed a diverse set of technical skills. "

        # Add work experience
        if work_experience:
            latest_job = work_experience[0]  # Get most recent experience
            experience_para += (f"In my current role as {latest_job.get('job_title')} at {latest_job.get('company_name')}, "
                              "I have successfully:")
            
            # Add bullet points from responsibilities
            responsibilities = latest_job.get('responsibilities', [])
            if responsibilities:
                experience_para += "\n"
                for resp in responsibilities[:3]:  # Include top 3 responsibilities
                    experience_para += f"\n• {resp}"

        # Add education if available
        if education:
            latest_edu = education[0]  # Get most recent education
            edu_para = f"\n\nI hold a {latest_edu.get('degree')} from {latest_edu.get('institution')}"
            if latest_edu.get('grade'):
                edu_para += f" with a grade of {latest_edu.get('grade')}"
            edu_para += ". "
            experience_para += edu_para

        # Create a closing paragraph
        closing = ("\n\nI am particularly drawn to your company's reputation for innovation and excellence in the industry. "
                  "I am excited about the possibility of bringing my expertise and enthusiasm to your team. "
                  "I would welcome the opportunity to discuss how my skills and experience align with your needs in more detail.")

        # Combine all paragraphs
        body = intro + experience_para + closing

        return body

    except Exception as e:
        # Fallback content in case of any error
        return (f"I am writing to express my strong interest in the {job_title} position at {company_name}.\n\n"
                "Throughout my career, I have developed a strong skill set that aligns well with this role. "
                "I have consistently demonstrated my ability to deliver results and work effectively in team environments. "
                "My technical expertise and problem-solving abilities have enabled me to make significant contributions "
                "in my previous roles.\n\n"
                "I am excited about the opportunity to bring my skills and experience to your organization. "
                "I would welcome the chance to discuss how I can contribute to your team in more detail.") 

def parse_cv_content(text: str) -> Dict:
    """Parse CV content into structured format using advanced pattern matching."""
    # Clean the input text first
    text = clean_text(text)
    
    content = {
        'personal_info': {
            'full_name': '',
            'email': '',
            'phone_number': '',
            'address': '',
            'nationality': '',
            'date_of_birth': '',
            'gender': 'M'  # Default value
        },
        'personal_profile': '',
        'work_experience': [],
        'education': [],
        'skills': [],
        'languages': [],
        'certifications': [],
        'hobbies': [],
        'referees': []
    }
    
    # Extract personal info using regex patterns
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_patterns = [
        r'(?:\+?\d{1,3}[-. ]?)?\(?\d{3}\)?[-. ]?\d{3}[-. ]?\d{4}',  # Standard format
        r'\+?\d{1,3}[-. ]?\d{3,4}[-. ]?\d{3,4}',  # International format
        r'\d{3}[-. ]?\d{3}[-. ]?\d{4}'  # Simple format
    ]
    name_patterns = [
        r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',  # Standard format
        r'^([A-Z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',  # ALL CAPS first name
        r'^((?:[A-Z][a-z]+\.?\s+)*[A-Z][a-z]+)'  # With middle initials
    ]
    address_patterns = [
        r'\b\d+\s+[A-Za-z\s,]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Way)[,\s]+[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}\b',
        r'\b\d+\s+[A-Za-z\s,]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Way)[,\s]+[A-Za-z\s]+\b',
        r'\b[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}\b'
    ]
    
    # Find email
    email_matches = re.finditer(email_pattern, text)
    for match in email_matches:
        content['personal_info']['email'] = match.group()
        break
    
    # Find phone number
    for pattern in phone_patterns:
        phone_matches = re.finditer(pattern, text)
        for match in phone_matches:
            phone = match.group()
            # Clean up phone number format
            phone = re.sub(r'[^\d+]', '', phone)  # Remove non-digit chars except +
            if len(phone) >= 10:  # Valid phone numbers should have at least 10 digits
                content['personal_info']['phone_number'] = phone
                break
        if content['personal_info']['phone_number']:
            break
    
    # Find full name (usually at the start of the CV)
    lines = text.split('\n')
    for line in lines[:10]:  # Check first 10 lines
        line = line.strip()
        for pattern in name_patterns:
            match = re.match(pattern, line)
            if match:
                name = match.group(1)
                # Verify it's not a section header
                if not any(header.lower() in name.lower() for header in SECTION_HEADERS):
                    content['personal_info']['full_name'] = name
                    break
        if content['personal_info']['full_name']:
            break
    
    # Find address
    for pattern in address_patterns:
        address_matches = re.finditer(pattern, text)
        for match in address_matches:
            content['personal_info']['address'] = match.group().strip()
            break
        if content['personal_info']['address']:
            break
    
    # Extract sections
    sections_to_extract = [
        ('personal_profile', ['profile', 'summary', 'about me', 'professional summary',
                          'career objective', 'objective'], 500),
        ('work_experience', ['work experience', 'employment history', 'professional experience',
                         'work history'], 2000),
        ('education', ['education', 'academic background', 'academic qualifications',
                    'educational background'], 1000),
        ('skills', ['skills', 'technical skills', 'key skills',
                 'professional skills', 'competencies'], 500),
        ('languages', ['languages', 'language skills', 'language proficiency'], 300),
        ('certifications', ['certifications', 'certificates', 'qualifications'], 500),
        ('hobbies', ['hobbies', 'interests', 'activities'], 300),
        ('referees', ['references', 'referees'], 500)
    ]
    
    for section_key, headers, max_chars in sections_to_extract:
        section_text = extract_section(text, headers, max_chars)
        if section_text:
            if section_key == 'work_experience':
                # Split into individual experiences
                experiences = []
                current_exp = []
                lines = section_text.split('\n')
                
                for line in lines:
                    # Check if this line might be the start of a new experience
                    if (re.search(r'\b\d{4}\b', line) or  # Contains a year
                        any(marker in line for marker in ['•', '-', '●', '○', '■', '▪', '►']) or  # Contains a bullet point
                        re.search(r'\b(?:at|@)\b', line)):  # Contains "at" or "@"
                        if current_exp:
                            exp_text = '\n'.join(current_exp)
                            parsed_exp = parse_work_experience(exp_text)
                            if parsed_exp:
                                experiences.append(parsed_exp)
                            current_exp = []
                    current_exp.append(line)
                
                # Don't forget the last experience
                if current_exp:
                    exp_text = '\n'.join(current_exp)
                    parsed_exp = parse_work_experience(exp_text)
                    if parsed_exp:
                        experiences.append(parsed_exp)
                
                content['work_experience'] = experiences
                
            elif section_key == 'education':
                # Split into individual education entries
                educations = []
                current_edu = []
                lines = section_text.split('\n')
                
                for line in lines:
                    # Check if this line might be the start of a new education entry
                    if (re.search(r'\b\d{4}\b', line) or  # Contains a year
                        any(marker in line for marker in ['•', '-', '●', '○', '■', '▪', '►']) or  # Contains a bullet point
                        any(term in line.lower() for term in ['university', 'college', 'school', 'institute', 'bachelor', 'master', 'phd', 'diploma'])):
                        if current_edu:
                            edu_text = '\n'.join(current_edu)
                            parsed_edu = parse_education(edu_text)
                            if parsed_edu:
                                educations.append(parsed_edu)
                            current_edu = []
                    current_edu.append(line)
                
                # Don't forget the last education entry
                if current_edu:
                    edu_text = '\n'.join(current_edu)
                    parsed_edu = parse_education(edu_text)
                    if parsed_edu:
                        educations.append(parsed_edu)
                
                content['education'] = educations
                
            elif section_key == 'skills':
                content['skills'] = parse_skills(section_text)
            elif section_key == 'languages':
                content['languages'] = parse_languages(section_text)
            elif section_key in ['personal_profile', 'hobbies']:
                # Clean and format the text
                cleaned_text = clean_text(section_text)
                # Remove any bullet points or section headers
                cleaned_text = re.sub(r'^[•\-●○■▪►]\s*', '', cleaned_text, flags=re.MULTILINE)
                content[section_key] = cleaned_text
    
    # Log parsing results
    logger.info(f"Parsed CV content - Found {len(content['work_experience'])} work experiences, "
                f"{len(content['education'])} education entries, "
                f"{len(content['skills'])} skills, and "
                f"{len(content['languages'])} languages")
    
    if not content['personal_info']['email'] or not content['personal_info']['phone_number']:
        logger.warning("Could not find email or phone number in CV content")
    
    return content

def extract_section(text, possible_headers, max_chars):
    """Extract a section from text based on possible headers."""
    text_lower = text.lower()
    for header in possible_headers:
        pattern = fr'\b{re.escape(header)}\b.*?(?=\n\s*\b(?:{'|'.join(SECTION_HEADERS)})\b|\Z)'
        match = re.search(pattern, text_lower, re.IGNORECASE | re.DOTALL)
        if match:
            section = match.group(0)
            if len(section) > max_chars:
                section = section[:max_chars]
            return section
    return ''

def parse_work_experience(text):
    """Parse a work experience entry."""
    lines = text.split('\n')
    job = {
        'job_title': '',
        'company_name': '',
        'location': '',
        'start_date': '',
        'end_date': '',
        'responsibilities': []
    }
    
    # Try to find job title and company
    for line in lines[:3]:  # Usually in first few lines
        if ' at ' in line.lower():
            parts = line.split(' at ')
            job['job_title'] = parts[0].strip()
            job['company_name'] = parts[1].strip()
            break
        elif ' - ' in line:
            parts = line.split(' - ')
            job['job_title'] = parts[0].strip()
            job['company_name'] = parts[1].strip()
            break
    
    # Try to find dates
    date_pattern = r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}'
    dates = re.findall(date_pattern, text)
    if len(dates) >= 2:
        job['start_date'] = dates[0]
        job['end_date'] = dates[1]
    elif len(dates) == 1:
        job['start_date'] = dates[0]
        job['end_date'] = 'Present'
    
    # Extract responsibilities
    resp_markers = ['•', '-', '●', '○', '■', '▪', '►']
    for line in lines:
        line = line.strip()
        if any(line.startswith(marker) for marker in resp_markers):
            resp = line.lstrip(''.join(resp_markers)).strip()
            if resp:
                job['responsibilities'].append(resp)
    
    return job if job['job_title'] and job['company_name'] else None

def parse_education(text):
    """Parse an education entry."""
    lines = text.split('\n')
    education = {
        'degree': '',
        'institution': '',
        'location': '',
        'start_date': '',
        'end_date': '',
        'grade': ''
    }
    
    # Try to find degree and institution
    for line in lines[:3]:
        if ' - ' in line:
            parts = line.split(' - ')
            education['degree'] = parts[0].strip()
            education['institution'] = parts[1].strip()
            break
        elif ',' in line:
            parts = line.split(',')
            education['degree'] = parts[0].strip()
            if len(parts) > 1:
                education['institution'] = parts[1].strip()
    
    # Try to find dates
    date_pattern = r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}'
    dates = re.findall(date_pattern, text)
    if len(dates) >= 2:
        education['start_date'] = dates[0]
        education['end_date'] = dates[1]
    elif len(dates) == 1:
        education['start_date'] = dates[0]
        education['end_date'] = 'Present'
    
    # Try to find grade
    grade_patterns = [
        r'Grade:\s*([\w\s.-]+)',
        r'GPA:\s*([\d.]+)',
        r'Class:\s*([\w\s.-]+)'
    ]
    for pattern in grade_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            education['grade'] = match.group(1).strip()
            break
    
    return education if education['degree'] and education['institution'] else None

def parse_skills(text):
    """Parse skills section into structured format."""
    skills = []
    skill_markers = ['•', '-', '●', '○', '■', '▪', '►', ',', '|']
    
    # First try to split by markers
    for marker in skill_markers:
        if marker in text:
            skill_list = [s.strip() for s in text.split(marker) if s.strip()]
            if len(skill_list) > 1:  # Found multiple skills
                return [{'name': skill, 'type': 'hard'} for skill in skill_list]
    
    # If no markers found, try line by line
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if line and not any(header.lower() in line.lower() for header in SECTION_HEADERS):
            skills.append({'name': line, 'type': 'hard'})
    
    return skills

def parse_languages(text):
    """Parse languages section into structured format."""
    languages = []
    lines = text.split('\n')
    
    proficiency_patterns = {
        r'\b(?:native|fluent|advanced|intermediate|basic)\b': lambda x: x.lower(),
        r'\b[cC][1-2]\b': 'fluent',
        r'\b[bB][1-2]\b': 'intermediate',
        r'\b[aA][1-2]\b': 'basic'
    }
    
    for line in lines:
        line = line.strip()
        if line:
            language = {'language': '', 'proficiency': 'basic'}
            
            # Try to find proficiency level
            for pattern, level in proficiency_patterns.items():
                match = re.search(pattern, line)
                if match:
                    prof = match.group(0)
                    language['proficiency'] = level(prof) if callable(level) else level
                    # Remove the proficiency part from the line
                    line = re.sub(pattern, '', line).strip()
                    break
            
            # The remaining text is the language name
            language['language'] = line.strip(' -•●○■▪►,')
            if language['language']:
                languages.append(language)
    
    return languages

def clean_text(text):
    """Clean and normalize text."""
    # Remove extra whitespace
    text = ' '.join(text.split())
    # Remove common section headers
    for header in SECTION_HEADERS:
        text = re.sub(fr'\b{re.escape(header)}\b', '', text, flags=re.IGNORECASE)
    return text.strip()

def split_into_entries(text):
    """Split a section into individual entries."""
    # Common patterns that indicate new entries
    patterns = [
        r'\n\d{4}',  # Year
        r'\n(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}',  # Month Year
        r'\n[A-Z][^a-z]{2,}',  # Uppercase words
        r'\n\s*•',  # Bullet points
        r'\n\s*-',  # Dashes
    ]
    
    entries = []
    current_entry = []
    lines = text.split('\n')
    
    for line in lines:
        is_new_entry = any(re.search(pattern, '\n' + line) for pattern in patterns)
        if is_new_entry and current_entry:
            entries.append('\n'.join(current_entry))
            current_entry = []
        current_entry.append(line)
    
    if current_entry:
        entries.append('\n'.join(current_entry))
    
    return entries

# Common section headers used in CVs
SECTION_HEADERS = [
    'profile', 'summary', 'about me', 'objective',
    'work experience', 'employment history', 'professional experience',
    'education', 'academic background', 'qualifications',
    'skills', 'technical skills', 'competencies',
    'languages', 'certifications', 'achievements',
    'interests', 'hobbies', 'references'
]

def parse_cover_letter_content(text: str) -> Dict:
    """Parse cover letter content into structured format."""
    content = {
        'personal_info': {
            'full_name': '',
            'email': '',
            'phone': '',
            'address': ''
        },
        'company_info': {
            'name': '',
            'address': ''
        },
        'date': '',
        'body': '',
        'salutation': 'Dear Hiring Manager,',
        'closing': 'Best regards,'
    }

    # Extract personal info using regex patterns
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_patterns = [
        r'(?:\+?\d{1,3}[-. ]?)?\(?\d{3}\)?[-. ]?\d{3}[-. ]?\d{4}',  # Standard format
        r'\+?\d{1,3}[-. ]?\d{3,4}[-. ]?\d{3,4}',  # International format
        r'\d{3}[-. ]?\d{3}[-. ]?\d{4}'  # Simple format
    ]
    name_patterns = [
        r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',  # Standard format
        r'^([A-Z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',  # ALL CAPS first name
        r'^((?:[A-Z][a-z]+\.?\s+)*[A-Z][a-z]+)'  # With middle initials
    ]
    address_patterns = [
        r'\b\d+\s+[A-Za-z\s,]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Way)[,\s]+[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}\b',
        r'\b\d+\s+[A-Za-z\s,]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Way)[,\s]+[A-Za-z\s]+\b',
        r'\b[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}\b'
    ]
    date_pattern = r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}'
    
    # Find email
    email_matches = re.finditer(email_pattern, text)
    for match in email_matches:
        content['personal_info']['email'] = match.group()
        break
    
    # Find phone number
    for pattern in phone_patterns:
        phone_matches = re.finditer(pattern, text)
        for match in phone_matches:
            phone = match.group()
            # Clean up phone number format
            phone = re.sub(r'[^\d+]', '', phone)  # Remove non-digit chars except +
            if len(phone) >= 10:  # Valid phone numbers should have at least 10 digits
                content['personal_info']['phone'] = phone
                break
        if content['personal_info']['phone']:
            break
    
    # Find full name (usually at the start of the cover letter)
    lines = text.split('\n')
    for line in lines[:10]:  # Check first 10 lines
        line = line.strip()
        for pattern in name_patterns:
            match = re.match(pattern, line)
            if match:
                name = match.group(1)
                content['personal_info']['full_name'] = name
                break
        if content['personal_info']['full_name']:
            break
    
    # Find address
    for pattern in address_patterns:
        address_matches = re.finditer(pattern, text)
        for match in address_matches:
            content['personal_info']['address'] = match.group().strip()
            break
        if content['personal_info']['address']:
            break
    
    # Find date
    date_matches = re.finditer(date_pattern, text)
    for match in date_matches:
        content['date'] = match.group()
        break
    
    # Extract body content (everything between salutation and closing)
    body_pattern = r'Dear.*?:\s*(.*?)\s*(?:Sincerely|Best regards|Yours truly|Regards|Best wishes|Yours sincerely|Yours faithfully)'
    body_match = re.search(body_pattern, text, re.DOTALL | re.IGNORECASE)
    if body_match:
        body = body_match.group(1).strip()
        # Clean up the body text
        body = re.sub(r'\s+', ' ', body)  # Replace multiple spaces with single space
        body = re.sub(r'\n\s*\n', '\n\n', body)  # Normalize paragraph breaks
        content['body'] = body
    else:
        # Fallback: try to extract text between first few lines and last few lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) > 10:
            content['body'] = '\n'.join(lines[3:-3])  # Skip potential header/footer lines
    
    return content

